# Copyright (C) 2025 Intel Corporation
# SPDX-License-Identifier: Apache-2.0
"""Multi-format document loader.

Extracts plain text from PDF, DOCX, TXT, and Markdown files so that the
downstream LLM can reason over a uniform representation.
"""

import os

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt", ".md", ".markdown")


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        return fh.read()


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _read_docx(path: str) -> str:
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    # Pull text out of tables as well, since SOPs often use them.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(c for c in cells if c))
    return "\n".join(parts)


def extract_text(path: str) -> str:
    """Return the extracted plain text for a supported document.

    Raises ValueError for unsupported extensions.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text = _read_pdf(path)
    elif ext in (".docx", ".doc"):
        text = _read_docx(path)
    elif ext in (".txt", ".md", ".markdown"):
        text = _read_txt(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    text = text.strip()
    if not text:
        raise ValueError(f"No extractable text found in '{os.path.basename(path)}'.")
    return text
